#!/usr/bin/env python3
"""
Create Test DOCX Files with Track Changes and Comments

Generates real DOCX files for integration testing of format verification module.
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


def add_track_changes_insertion(paragraph, text, author="Test User"):
    """Add insertion (Track Changes) to paragraph."""
    run = paragraph.add_run(text)

    # Create insertion element
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), '0')
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), '2025-11-02T10:00:00Z')

    # Move run into insertion
    run._element.getparent().remove(run._element)
    ins.append(run._element)
    paragraph._element.append(ins)

    return run


def add_track_changes_deletion(paragraph, text, author="Test User"):
    """Add deletion (Track Changes) to paragraph."""
    run = paragraph.add_run(text)

    # Create deletion element
    dele = OxmlElement('w:del')
    dele.set(qn('w:id'), '1')
    dele.set(qn('w:author'), author)
    dele.set(qn('w:date'), '2025-11-02T10:00:00Z')

    # Move run into deletion
    run._element.getparent().remove(run._element)
    dele.append(run._element)
    paragraph._element.append(dele)

    return run


def create_docx_with_track_changes(output_path: Path):
    """Create DOCX with Track Changes (insertions and deletions)."""
    doc = Document()

    # Add title
    doc.add_heading('Test Document with Track Changes', 0)

    # Paragraph 1: Normal text
    p1 = doc.add_paragraph('This is normal text without any changes. ')

    # Paragraph 2: Text with insertion
    p2 = doc.add_paragraph('This sentence has ')
    add_track_changes_insertion(p2, 'inserted text')
    p2.add_run(' in the middle.')

    # Paragraph 3: Text with deletion
    p3 = doc.add_paragraph('This sentence has ')
    add_track_changes_deletion(p3, 'deleted text')
    p3.add_run(' marked for removal.')

    # Paragraph 4: Multiple changes
    p4 = doc.add_paragraph('Multiple changes: ')
    add_track_changes_insertion(p4, 'added word')
    p4.add_run(', ')
    add_track_changes_deletion(p4, 'removed word')
    p4.add_run(', and ')
    add_track_changes_insertion(p4, 'another insertion')
    p4.add_run('.')

    # Save
    doc.save(output_path)
    print(f"✅ Created: {output_path}")
    print(f"   - 3 insertions, 2 deletions = 5 Track Changes total")


def create_docx_without_track_changes(output_path: Path):
    """Create clean DOCX without Track Changes."""
    doc = Document()

    doc.add_heading('Clean Test Document', 0)
    doc.add_paragraph('This is normal text without any Track Changes.')
    doc.add_paragraph('No insertions or deletions marked.')
    doc.add_paragraph('Clean document for comparison testing.')

    doc.save(output_path)
    print(f"✅ Created: {output_path}")
    print(f"   - 0 Track Changes")


def create_docx_with_comments(output_path: Path):
    """Create DOCX with comments (simplified - requires comments.xml manipulation)."""
    # Note: python-docx doesn't fully support comments yet
    # This creates a basic document; comments would need manual addition or lxml manipulation
    doc = Document()

    doc.add_heading('Test Document for Comments', 0)
    doc.add_paragraph('This document should have comments added manually.')
    doc.add_paragraph('Comment testing requires comments.xml in DOCX structure.')

    doc.save(output_path)
    print(f"⚠️  Created: {output_path}")
    print(f"   - Comment support limited in python-docx")
    print(f"   - Manual comment addition recommended for testing")


def create_all_test_files():
    """Generate all test DOCX files."""
    fixtures_dir = Path(__file__).parent
    fixtures_dir.mkdir(exist_ok=True)

    print("Creating test DOCX files...")
    print("=" * 60)

    # Test file 1: With Track Changes
    create_docx_with_track_changes(
        fixtures_dir / "test_with_track_changes.docx"
    )

    # Test file 2: Clean (no Track Changes)
    create_docx_without_track_changes(
        fixtures_dir / "test_clean.docx"
    )

    # Test file 3: For comments (placeholder)
    create_docx_with_comments(
        fixtures_dir / "test_with_comments.docx"
    )

    print("=" * 60)
    print("✅ Test file generation complete")
    print(f"Location: {fixtures_dir}")


if __name__ == "__main__":
    try:
        create_all_test_files()
    except ImportError:
        print("❌ ERROR: python-docx not installed")
        print("Install: pip install python-docx")
        exit(1)
